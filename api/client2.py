from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import os
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import GPT4AllEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.messages import HumanMessage
from docx import Document as DocxDocument
import time
from pymongo import MongoClient
from dotenv import load_dotenv

app = FastAPI()

# Load environment variables from .env file
load_dotenv()

# MongoDB connection
mongodb_uri = os.getenv("MONGODB_URI")
client = MongoClient(mongodb_uri)
db = client.get_database("chat_db")
chat_collection = db.get_collection("chat_history")

# Initialize model and components
local_llm = 'llama3'
cascade_directory = "../data/cascade"
policy_directory = "../data/policy"
directories = [cascade_directory, policy_directory]
pdf_file_paths = []
docx_file_paths = []


# MongoDB connection
mongodb_uri = os.getenv("MONGODB_URI")
if not mongodb_uri:
    raise ValueError("No MONGODB_URI found in environment variables")

try:
    client = MongoClient(mongodb_uri)
    db = client.get_database("chat_db")
    chat_collection = db.get_collection("chat_history")
    # Test the connection
    client.admin.command('ping')
    print("MongoDB connection established successfully.")
except ConnectionError as e:
    print(f"MongoDB connection error: {e}")
    raise HTTPException(status_code=500, detail=str(e))


for directory in directories:
    for filename in os.listdir(directory):
        if filename.endswith(".pdf"):
            pdf_file_paths.append(os.path.join(directory, filename))
        elif filename.endswith(".docx"):
            docx_file_paths.append(os.path.join(directory, filename))

docs_list = []

for pdf_path in pdf_file_paths:
    loader = PyMuPDFLoader(pdf_path)
    try:
        loaded_docs = loader.load()
        docs_list.extend(loaded_docs)
    except Exception as e:
        print(f"Error loading {pdf_path}: {e}")

for docx_path in docx_file_paths:
    try:
        docx = DocxDocument(docx_path)
        full_text = []
        for para in docx.paragraphs:
            full_text.append(para.text)
        docx_text = '\n'.join(full_text)
        docs_list.append(Document(page_content=docx_text, metadata={"source": docx_path}))
    except Exception as e:
        print(f"Error loading {docx_path}: {e}")

text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(chunk_size=250, chunk_overlap=0)
doc_splitter = text_splitter.split_documents(docs_list)

filtered_doc = []
for doc in doc_splitter:
    if isinstance(doc, Document) and hasattr(doc, 'metadata'):
        clean_metadata = {k: v for k, v in doc.metadata.items() if isinstance(v, (str, int, float, bool))}
        filtered_doc.append(Document(page_content=doc.page_content, metadata=clean_metadata))

embedding = GPT4AllEmbeddings(model_name="all-MiniLM-L6-v2.gguf2.f16.gguf", gpt4all_kwargs={'allow_download': 'True'})
vectorstore = Chroma.from_documents(
    documents=filtered_doc,
    collection_name="rag-chroma",
    embedding=embedding,
)

retriever = vectorstore.as_retriever()
llm = ChatOllama(model=local_llm, temperature=0)

contextualize_q_system_prompt = """Given a chat history and the latest user question \
which might reference context in the chat history, formulate a standalone question \
which can be understood without the chat history. Do NOT answer the question, \
just reformulate it if needed and otherwise return it as is."""
contextualize_q_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ]
)
history_aware_retriever = create_history_aware_retriever(
    llm, retriever, contextualize_q_prompt
)

qa_system_prompt = """system You are an assistant for question-answering tasks. Use the following context to answer the question. Avoid phrases like "Based on the provided context". Explain the answer in the end. and make a heading with paragraph.
Question: {input}
Context: {context}
Answer: assistant"""
qa_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", qa_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ]
)

question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

chat_history = []

class QueryModel(BaseModel):
    question: str

class AnswerResponse(BaseModel):
    answer: str
    responseTime: str

def serialize_document(document):
    document["_id"] = str(document["_id"])
    return document

@app.post("/ask")
async def ask_question(query: QueryModel):
    start_time = time.time()
    try:
        question = query.question
        response = rag_chain.invoke({"input": question, "chat_history": chat_history})
        response_time = "{:.2f} sec".format(time.time() - start_time)

        # Convert HumanMessage to dict for storage
        human_message_dict = {"type": "human", "content": question}
        assistant_message_dict = {"type": "assistant", "content": response["answer"]}

        chat_history.extend([HumanMessage(content=question), response["answer"]])

        # New messages to be appended
        new_messages = [
            human_message_dict,
            assistant_message_dict
        ]

        # Attempt to update the document
        result = chat_collection.update_one(
            {"identifier": "231"},
            {
                "$set": {
                    "timestamp": time.time()
                },
                "$push": {
                    "messages":  new_messages
                }
            }
        )

        # If no document was matched and updated, insert a new document
        if result.matched_count == 0:
            new_document = {
                "identifier": "231",
                "messages": [new_messages],
                "user": "haris-ali",
                "promptTitle": "Prompt-Title",
                "timestamp": time.time()
            }
            chat_collection.insert_one(new_document)
            print("Document not found, inserted a new document.")
        else:
            print("Document updated successfully!")

        return AnswerResponse(answer=response["answer"], responseTime=response_time)

    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

# Example function to retrieve chat history from MongoDB

@app.get("/chat_history")
async def get_chat_history():
    try:
        history_docs = chat_collection.find({"identifier": "231"})
        history_list = [serialize_document(doc) for doc in history_docs]

        if history_list:
            messages = history_list[0]["messages"]

            for message_pair in messages:
                if len(message_pair) == 2:
                    human_message = message_pair[0]
                    assistant_message = message_pair[1]

                    if human_message["type"] == "human" and assistant_message["type"] == "assistant":
                        chat_history.extend([
                            HumanMessage(content=human_message["content"]),
                            assistant_message["content"]
                        ])

        return history_list
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/clear")
async def get_chat_history_clear():
    chat_history.clear()

@app.get("/get")
async def get_chat_history_get():
    return chat_history

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
